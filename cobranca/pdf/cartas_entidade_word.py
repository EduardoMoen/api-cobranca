import datetime

from io import BytesIO
from collections import defaultdict
from textwrap import dedent

from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def gerar_carta_entidade_word(dividas):
    document = Document()

    #MARGENS
    section = document.sections[0]

    section.page_width = Mm(210)
    section.page_height = Mm(297)

    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    #RODAPÉ
    footer = section.footer

    p_footer = footer.paragraphs[0]

    p_footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p_footer.add_run(
        "Rua Barbosa da Cunha, 386 - Guanabara, Campinas/SP - CEP 13.073-320\n"
        "Fone: (19) 3232-5767"
    )

    run.font.size = Pt(9)

    #TEXTOS
    texto2 = dedent("""Tendo esgotado o prazo para pagamento no Banco Credenciado, V.Sa. deverá, dentro do prazo máximo de 5 (cinco) dias, entrar em contato com o Escritório Jurídico, que estará atendendo de segunda à sexta-feira das 9:00 às 12:00 h e das 13:00 às 17:00 h, pelo Telefone/WhatsApp (19) 3232-5767.""")

    texto3 = dedent("""Esgotando este prazo, e no seu silêncio, serão tomadas providências cabíveis junto aos serviços de proteção ao crédito, e, posteriormente a propositura de ação judicial.""")

    texto4 = dedent("""Caso V. Sa. já tenha pago o seu débito, queira por gentileza, nos enviar comprovação do pagamento e desconsiderar esta comunicação.""")

    #AGRUPAR RESPONSAVEIS
    dividas_por_responsavel = defaultdict(list)

    for divida in dividas:
        dividas_por_responsavel[divida.responsavel].append(divida)

    total_responsaveis = len(dividas_por_responsavel)

    contador_responsavel = 0

    for responsavel, divs_resp in dividas_por_responsavel.items():
        contador_responsavel += 1

        dividas_por_aluno = defaultdict(list)

        for divida in divs_resp:
            dividas_por_aluno[divida.nomeAluno].append(divida)

        total_alunos = len(dividas_por_aluno)

        contador_aluno = 0

        for nome_aluno, lista_dividas in dividas_por_aluno.items():

            contador_aluno += 1

            numero_aluno = lista_dividas[0].codigoAluno
            nome_escola = lista_dividas[0].escola.nome

            texto1 = dedent(f"""Pela presente comunicamos V.Sa. que, encontra-se em aberto o pagamento da(s) mensalidade(s) de seu filho(a) matriculado(a) no(a) {nome_escola} referente(s) aos boleto(s) vencido(s) em:""")

            #LOGO
            try:
                document.add_picture(
                    "cobranca/pdf/logo/logo_carta.jpeg",
                    width=Inches(1.2)
                )
            except:
                pass

            #DATA
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.add_run(
                f"Campinas, {datetime.date.today().strftime('%d/%m/%Y')}."
            )

            document.add_paragraph()

            p = document.add_paragraph()
            p.add_run("Ilmo(a). Sr(a).\n")
            p.add_run(f"{responsavel.nome} - {responsavel.cpf}\n")
            p.add_run(
                f"Aluno: {nome_aluno} - {numero_aluno}\n"
            )
            p.add_run(
                f"{responsavel.endereco} - {responsavel.bairro}\n"
            )
            p.add_run(
                f"{responsavel.cep} {responsavel.cidade} - {responsavel.uf}"
            )

            document.add_paragraph()

            p = document.add_paragraph(texto1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            document.add_paragraph()

            boletos = ", ".join([
                f"{divida.numeroCobranca}-{divida.dataVencimento.month}-{divida.dataVencimento.year}"
                for divida in lista_dividas
            ])

            document.add_paragraph(boletos)

            document.add_paragraph()

            p = document.add_paragraph(texto2)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            document.add_paragraph()

            p = document.add_paragraph(texto3)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            document.add_paragraph()

            p=document.add_paragraph(texto4)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            document.add_paragraph()

            document.add_paragraph("Atenciosamente,")

            document.add_paragraph()

            document.add_paragraph("Cremovale Cobranças e Serviços Ltda")

            #QUEBRA ENTRE ALUNOS
            if contador_aluno < total_alunos:
                document.add_page_break()

        #QUEBRA ENTRE RESPONSAVEIS
        if contador_responsavel < total_responsaveis:
            document.add_page_break()

    #SALVAR

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer
